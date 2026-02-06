"""报告导出服务"""
import pandas as pd
from typing import Dict, Any, List, Optional, BinaryIO
from datetime import datetime
from pathlib import Path
import base64
from io import BytesIO


class ReportService:
    """报告导出服务类"""
    
    @staticmethod
    def _format_number(value: float, decimals: int = 2) -> str:
        """格式化数字"""
        if value is None:
            return "N/A"
        return f"{value:,.{decimals}f}"
    
    @staticmethod
    def _generate_html_report(
        title: str,
        dataset_info: Dict[str, Any],
        analysis_results: List[Dict[str, Any]],
        ai_summary: str = None,
        generated_at: str = None
    ) -> str:
        """生成HTML格式报告"""
        if not generated_at:
            generated_at = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        
        html = f"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
    <meta charset="UTF-8">
    <title>{title}</title>
    <style>
        * {{ margin: 0; padding: 0; box-sizing: border-box; }}
        body {{
            font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", Roboto, "Helvetica Neue", Arial, sans-serif;
            line-height: 1.6;
            color: #333;
            max-width: 1200px;
            margin: 0 auto;
            padding: 40px;
            background: #f5f5f5;
        }}
        .container {{
            background: white;
            padding: 40px;
            border-radius: 8px;
            box-shadow: 0 2px 8px rgba(0,0,0,0.1);
        }}
        .header {{
            text-align: center;
            padding-bottom: 30px;
            border-bottom: 3px solid #3b82f6;
            margin-bottom: 30px;
        }}
        .header h1 {{
            color: #1e40af;
            font-size: 32px;
            margin-bottom: 10px;
        }}
        .meta {{
            color: #6b7280;
            font-size: 14px;
        }}
        .section {{
            margin: 30px 0;
            padding: 20px;
            background: #f8fafc;
            border-radius: 8px;
            border-left: 4px solid #3b82f6;
        }}
        .section h2 {{
            color: #1e40af;
            font-size: 22px;
            margin-bottom: 15px;
            padding-bottom: 10px;
            border-bottom: 1px solid #e5e7eb;
        }}
        .section h3 {{
            color: #374151;
            font-size: 18px;
            margin: 20px 0 10px 0;
        }}
        .info-grid {{
            display: grid;
            grid-template-columns: repeat(auto-fit, minmax(200px, 1fr));
            gap: 15px;
            margin: 20px 0;
        }}
        .info-card {{
            background: white;
            padding: 15px;
            border-radius: 6px;
            border: 1px solid #e5e7eb;
        }}
        .info-card .label {{
            color: #6b7280;
            font-size: 12px;
            text-transform: uppercase;
            margin-bottom: 5px;
        }}
        .info-card .value {{
            color: #1f2937;
            font-size: 20px;
            font-weight: 600;
        }}
        table {{
            width: 100%;
            border-collapse: collapse;
            margin: 15px 0;
            background: white;
        }}
        th, td {{
            padding: 12px;
            text-align: left;
            border-bottom: 1px solid #e5e7eb;
        }}
        th {{
            background: #f3f4f6;
            font-weight: 600;
            color: #374151;
        }}
        tr:hover {{
            background: #f9fafb;
        }}
        .stats-table td {{
            font-family: 'Courier New', monospace;
        }}
        .chart-container {{
            text-align: center;
            margin: 20px 0;
            padding: 20px;
            background: white;
            border-radius: 8px;
            border: 1px solid #e5e7eb;
        }}
        .chart-container img {{
            max-width: 100%;
            height: auto;
            border-radius: 4px;
        }}
        .chart-title {{
            font-weight: 600;
            color: #374151;
            margin-bottom: 10px;
        }}
        .insight-box {{
            background: #eff6ff;
            border: 1px solid #bfdbfe;
            border-radius: 8px;
            padding: 20px;
            margin: 15px 0;
        }}
        .insight-box h4 {{
            color: #1e40af;
            margin-bottom: 10px;
        }}
        .footer {{
            text-align: center;
            margin-top: 40px;
            padding-top: 20px;
            border-top: 1px solid #e5e7eb;
            color: #9ca3af;
            font-size: 12px;
        }}
        .badge {{
            display: inline-block;
            padding: 4px 8px;
            border-radius: 4px;
            font-size: 12px;
            font-weight: 500;
        }}
        .badge-success {{ background: #dcfce7; color: #166534; }}
        .badge-warning {{ background: #fef3c7; color: #92400e; }}
        .badge-info {{ background: #dbeafe; color: #1e40af; }}
    </style>
</head>
<body>
    <div class="container">
        <div class="header">
            <h1>📊 {title}</h1>
            <div class="meta">
                生成时间：{generated_at}<br>
                InsightEase 数据分析平台
            </div>
        </div>
"""
        
        # 数据集信息
        html += f"""
        <div class="section">
            <h2>📁 数据集信息</h2>
            <div class="info-grid">
                <div class="info-card">
                    <div class="label">文件名</div>
                    <div class="value">{dataset_info.get('filename', 'N/A')}</div>
                </div>
                <div class="info-card">
                    <div class="label">数据行数</div>
                    <div class="value">{dataset_info.get('row_count', 0):,}</div>
                </div>
                <div class="info-card">
                    <div class="label">数据列数</div>
                    <div class="value">{dataset_info.get('col_count', 0)}</div>
                </div>
                <div class="info-card">
                    <div class="label">文件大小</div>
                    <div class="value">{ReportService._format_file_size(dataset_info.get('file_size', 0))}</div>
                </div>
            </div>
"""
        
        # 列信息
        schema = dataset_info.get('schema', [])
        if schema:
            html += """
            <h3>数据列结构</h3>
            <table>
                <thead>
                    <tr>
                        <th>列名</th>
                        <th>数据类型</th>
                        <th>示例值</th>
                    </tr>
                </thead>
                <tbody>
"""
            for col in schema[:20]:  # 最多显示20列
                sample = col.get('sample_values', [])
                sample_str = ', '.join(str(s) for s in sample[:3]) if sample else '-'
                html += f"""
                    <tr>
                        <td>{col.get('name', '-')}</td>
                        <td><span class="badge badge-info">{col.get('dtype', '-')}</span></td>
                        <td>{sample_str}</td>
                    </tr>
"""
            html += """
                </tbody>
            </table>
"""
        
        html += "</div>"
        
        # AI 摘要
        if ai_summary:
            html += f"""
        <div class="section">
            <h2>🤖 AI 数据洞察</h2>
            <div class="insight-box">
                {ai_summary}
            </div>
        </div>
"""
        
        # 分析结果
        if analysis_results:
            for idx, result in enumerate(analysis_results, 1):
                analysis_type = result.get('type', '')
                result_data = result.get('data', {})
                
                html += f"""
        <div class="section">
            <h2>📈 分析 {idx}：{ReportService._get_analysis_type_name(analysis_type)}</h2>
"""
                
                # 根据分析类型渲染不同内容
                if analysis_type == 'descriptive':
                    html += ReportService._render_descriptive_stats(result_data)
                elif analysis_type == 'correlation':
                    html += ReportService._render_correlation(result_data)
                elif analysis_type == 'visualization':
                    html += ReportService._render_visualization(result_data)
                elif analysis_type == 'forecast':
                    html += ReportService._render_forecast(result_data)
                else:
                    # 通用渲染
                    html += f"<pre>{ReportService._format_dict(result_data)}</pre>"
                
                html += "</div>"
        
        # 页脚
        html += """
        <div class="footer">
            <p>本报告由 InsightEase 数据分析平台自动生成</p>
            <p>© 2024 InsightEase. All rights reserved.</p>
        </div>
    </div>
</body>
</html>
"""
        
        return html
    
    @staticmethod
    def _get_analysis_type_name(analysis_type: str) -> str:
        """获取分析类型中文名"""
        names = {
            'descriptive': '描述性统计分析',
            'correlation': '相关性分析',
            'distribution': '分布分析',
            'outlier': '异常值检测',
            'visualization': '数据可视化',
            'forecast': '趋势预测',
            'comprehensive': '综合分析'
        }
        return names.get(analysis_type, analysis_type)
    
    @staticmethod
    def _format_file_size(size_bytes: int) -> str:
        """格式化文件大小"""
        if size_bytes < 1024:
            return f"{size_bytes} B"
        elif size_bytes < 1024 * 1024:
            return f"{size_bytes / 1024:.1f} KB"
        else:
            return f"{size_bytes / (1024 * 1024):.1f} MB"
    
    @staticmethod
    def _render_descriptive_stats(data: Dict[str, Any]) -> str:
        """渲染描述性统计"""
        html = ""
        column_stats = data.get('column_stats', [])
        
        if column_stats:
            html += "<h3>数值型列统计</h3><table class='stats-table'>"
            html += "<tr><th>列名</th><th>均值</th><th>中位数</th><th>标准差</th><th>最小值</th><th>最大值</th></tr>"
            
            for col in column_stats:
                if col.get('type') == 'numeric':
                    html += f"""
                    <tr>
                        <td>{col.get('name', '-')}</td>
                        <td>{ReportService._format_number(col.get('mean'))}</td>
                        <td>{ReportService._format_number(col.get('median'))}</td>
                        <td>{ReportService._format_number(col.get('std'))}</td>
                        <td>{ReportService._format_number(col.get('min'))}</td>
                        <td>{ReportService._format_number(col.get('max'))}</td>
                    </tr>
"""
            html += "</table>"
        
        return html
    
    @staticmethod
    def _render_correlation(data: Dict[str, Any]) -> str:
        """渲染相关性分析"""
        html = ""
        strong_corrs = data.get('strong_correlations', [])
        
        if strong_correlations:
            html += "<h3>强相关性发现</h3><table>"
            html += "<tr><th>列1</th><th>列2</th><th>相关系数</th><th>关系</th></tr>"
            
            for corr in strong_corrs:
                html += f"""
                <tr>
                    <td>{corr.get('column1', '-')}</td>
                    <td>{corr.get('column2', '-')}</td>
                    <td>{ReportService._format_number(corr.get('correlation'))}</td>
                    <td>{corr.get('strength', '-')}</td>
                </tr>
"""
            html += "</table>"
        else:
            html += "<p>未发现强相关性（|r| > 0.7）</p>"
        
        return html
    
    @staticmethod
    def _render_visualization(data: Dict[str, Any]) -> str:
        """渲染可视化图表"""
        html = ""
        charts = data.get('charts', [])
        
        for chart in charts:
            if 'image_base64' in chart:
                html += f"""
                <div class="chart-container">
                    <div class="chart-title">{chart.get('type', 'Chart')}</div>
                    <img src="data:image/png;base64,{chart['image_base64']}" alt="chart">
                </div>
"""
        
        return html
    
    @staticmethod
    def _render_forecast(data: Dict[str, Any]) -> str:
        """渲染预测结果"""
        html = ""
        
        if 'statistics' in data:
            stats = data['statistics']
            html += """
            <div class="info-grid">
                <div class="info-card">
                    <div class="label">历史均值</div>
                    <div class="value">""" + ReportService._format_number(stats.get('historical_mean')) + """</div>
                </div>
                <div class="info-card">
                    <div class="label">预测均值</div>
                    <div class="value">""" + ReportService._format_number(stats.get('forecast_mean')) + """</div>
                </div>
                <div class="info-card">
                    <div class="label">趋势方向</div>
                    <div class="value">""" + stats.get('trend_direction', '-') + """</div>
                </div>
            </div>
"""
        
        return html
    
    @staticmethod
    def _format_dict(data: Dict[str, Any], indent: int = 0) -> str:
        """格式化字典为可读文本"""
        lines = []
        for key, value in data.items():
            if isinstance(value, dict):
                lines.append(f"{'  ' * indent}{key}:")
                lines.append(ReportService._format_dict(value, indent + 1))
            elif isinstance(value, list):
                lines.append(f"{'  ' * indent}{key}: [{len(value)} items]")
            else:
                lines.append(f"{'  ' * indent}{key}: {value}")
        return '\n'.join(lines)
    
    @staticmethod
    def generate_pdf_report(
        title: str,
        dataset_info: Dict[str, Any],
        analysis_results: List[Dict[str, Any]],
        ai_summary: str = None
    ) -> bytes:
        """
        生成PDF报告
        """
        try:
            from weasyprint import HTML, CSS
            
            # 生成HTML
            html_content = ReportService._generate_html_report(
                title, dataset_info, analysis_results, ai_summary
            )
            
            # 转换为PDF
            html = HTML(string=html_content)
            pdf_bytes = html.write_pdf()
            
            return pdf_bytes
            
        except ImportError:
            # 如果没有weasyprint，返回错误提示
            raise Exception("PDF生成功能需要安装weasyprint: pip install weasyprint")
        except Exception as e:
            raise Exception(f"PDF生成失败: {str(e)}")
    
    @staticmethod
    def generate_word_report(
        title: str,
        dataset_info: Dict[str, Any],
        analysis_results: List[Dict[str, Any]],
        ai_summary: str = None
    ) -> bytes:
        """
        生成Word报告
        """
        try:
            from docx import Document
            from docx.shared import Inches, Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            
            doc = Document()
            
            # 标题
            title_heading = doc.add_heading(title, 0)
            title_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 元信息
            meta = doc.add_paragraph()
            meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
            meta.add_run(f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n").italic = True
            meta.add_run("InsightEase 数据分析平台").italic = True
            
            doc.add_paragraph()
            
            # 数据集信息
            doc.add_heading('📁 数据集信息', level=1)
            
            info_table = doc.add_table(rows=2, cols=4)
            info_table.style = 'Light Grid Accent 1'
            
            headers = ['文件名', '数据行数', '数据列数', '文件大小']
            values = [
                dataset_info.get('filename', 'N/A'),
                f"{dataset_info.get('row_count', 0):,}",
                str(dataset_info.get('col_count', 0)),
                ReportService._format_file_size(dataset_info.get('file_size', 0))
            ]
            
            for i, (header, value) in enumerate(zip(headers, values)):
                info_table.rows[0].cells[i].text = header
                info_table.rows[1].cells[i].text = value
            
            doc.add_paragraph()
            
            # 列结构
            schema = dataset_info.get('schema', [])
            if schema:
                doc.add_heading('数据列结构', level=2)
                
                col_table = doc.add_table(rows=1, cols=3)
                col_table.style = 'Light Grid Accent 1'
                
                hdr_cells = col_table.rows[0].cells
                hdr_cells[0].text = '列名'
                hdr_cells[1].text = '数据类型'
                hdr_cells[2].text = '示例值'
                
                for col in schema[:20]:
                    row_cells = col_table.add_row().cells
                    row_cells[0].text = col.get('name', '-')
                    row_cells[1].text = col.get('dtype', '-')
                    sample = col.get('sample_values', [])
                    row_cells[2].text = ', '.join(str(s) for s in sample[:3]) if sample else '-'
                
                doc.add_paragraph()
            
            # AI 摘要
            if ai_summary:
                doc.add_heading('🤖 AI 数据洞察', level=1)
                doc.add_paragraph(ai_summary)
                doc.add_paragraph()
            
            # 分析结果
            if analysis_results:
                doc.add_heading('📈 分析结果', level=1)
                
                for idx, result in enumerate(analysis_results, 1):
                    analysis_type = result.get('type', '')
                    result_data = result.get('data', {})
                    
                    doc.add_heading(f'分析 {idx}：{ReportService._get_analysis_type_name(analysis_type)}', level=2)
                    
                    # 添加结果文本
                    result_text = ReportService._format_dict(result_data)
                    doc.add_paragraph(result_text[:2000])  # 限制长度
                    doc.add_paragraph()
            
            # 页脚
            doc.add_paragraph()
            footer = doc.add_paragraph()
            footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
            footer_run = footer.add_run('本报告由 InsightEase 数据分析平台自动生成\n© 2024 InsightEase. All rights reserved.')
            footer_run.font.size = Pt(9)
            footer_run.font.color.rgb = RGBColor(128, 128, 128)
            
            # 保存到内存
            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            
            return buffer.getvalue()
            
        except ImportError:
            raise Exception("Word生成功能需要安装python-docx: pip install python-docx")
        except Exception as e:
            raise Exception(f"Word生成失败: {str(e)}")
    
    @staticmethod
    def generate_html_file(
        title: str,
        dataset_info: Dict[str, Any],
        analysis_results: List[Dict[str, Any]],
        ai_summary: str = None
    ) -> str:
        """
        生成HTML文件内容
        """
        return ReportService._generate_html_report(
            title, dataset_info, analysis_results, ai_summary
        )


# 单例实例
report_service = ReportService()
